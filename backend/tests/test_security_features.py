import io
import zipfile

import pytest
from docx import Document

from app.agents.law_agent import _enforce_source_policy
from app.agents.response_standard import append_legal_standard
from app.core.security import totp_code
from app.utils.document_parser import parse_file, pdf_needs_ocr
from app.utils.upload_security import UploadSecurityError, validate_file_signature


async def test_cookie_refresh_logout_and_mfa(client):
    payload = {
        "email": "security@test.uz",
        "username": "security-user",
        "password": "SecurePassword2026",
    }
    assert (await client.post("/api/auth/register", json=payload)).status_code == 201

    login = await client.post(
        "/api/auth/login",
        json={"email": payload["email"], "password": payload["password"]},
    )
    assert login.status_code == 200
    assert "auth_token" in client.cookies
    assert "refresh_token" in client.cookies
    assert (await client.get("/api/auth/me")).status_code == 200

    setup = await client.post("/api/auth/mfa/setup")
    assert setup.status_code == 200
    code = totp_code(setup.json()["secret"])
    assert (await client.post("/api/auth/mfa/enable", json={"code": code})).status_code == 200

    await client.post("/api/auth/logout")
    assert (await client.get("/api/auth/me")).status_code == 401
    missing_mfa = await client.post(
        "/api/auth/login",
        json={"email": payload["email"], "password": payload["password"]},
    )
    assert missing_mfa.status_code == 403
    with_mfa = await client.post(
        "/api/auth/login",
        json={
            "email": payload["email"],
            "password": payload["password"],
            "mfa_code": totp_code(setup.json()["secret"]),
        },
    )
    assert with_mfa.status_code == 200
    old_refresh = client.cookies.get("refresh_token")
    assert (await client.post("/api/auth/refresh")).status_code == 200
    assert client.cookies.get("refresh_token") != old_refresh


async def test_weak_password_is_rejected(client):
    response = await client.post(
        "/api/auth/register",
        json={"email": "weak@test.uz", "username": "weak-user", "password": "abcdefghij"},
    )
    assert response.status_code == 422


def test_upload_signature_and_archive_limits(monkeypatch):
    with pytest.raises(UploadSecurityError):
        validate_file_signature("fake.pdf", b"not a pdf")

    archive_bytes = io.BytesIO()
    monkeypatch.setattr("app.utils.upload_security.MAX_UNCOMPRESSED_SIZE", 10)
    with zipfile.ZipFile(archive_bytes, "w") as archive:
        archive.writestr("[Content_Types].xml", "types")
        archive.writestr("word/document.xml", "x" * 11)
    with pytest.raises(UploadSecurityError):
        validate_file_signature("bomb.docx", archive_bytes.getvalue())

    with pytest.raises(ValueError):
        parse_file("broken.pdf", b"%PDF-1.7\ninvalid")


def test_docx_tables_are_included_and_empty_pdf_needs_ocr():
    document = Document()
    document.add_paragraph("Основные условия")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Обязательство"
    table.cell(0, 1).text = "Срок"
    table.cell(1, 0).text = "Оплатить"
    table.cell(1, 1).text = "10 дней"
    payload = io.BytesIO()
    document.save(payload)

    text = parse_file("obligations.docx", payload.getvalue())
    assert "Основные условия" in text
    assert "Обязательство | Срок" in text
    assert "Оплатить | 10 дней" in text
    assert pdf_needs_ocr("scan.pdf", "") is True
    assert pdf_needs_ocr("searchable.pdf", "Договор " * 40) is False


def test_unverified_legal_citation_is_removed():
    sources = [
        {
            "document_title": "Verified act",
            "article_number": "21",
            "url": "https://lex.uz/ru/docs/10872#11670",
        }
    ]
    result = _enforce_source_policy(
        {
            "legal_issues": [
                {
                    "issue": "Claim",
                    "article": "article 999",
                    "source_url": "https://lex.uz/ru/docs/fabricated",
                }
            ]
        },
        sources,
    )
    issue = result["legal_issues"][0]
    assert issue["citation_verified"] is False
    assert issue["article"] is None
    assert issue["source_url"] is None

    text = append_legal_standard(
        "Fabricated: https://lex.uz/ru/docs/fabricated",
        sources,
    )
    assert "https://lex.uz/ru/docs/fabricated" not in text
    assert "https://lex.uz/ru/docs/10872#11670" in text

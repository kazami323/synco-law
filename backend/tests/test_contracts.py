"""Тесты контрактов: CRUD, версии, файлы, видимость по ролям, метрики."""

import io

from docx import Document

from tests.conftest import register_and_login

CONTRACT = {
    "title": "Договор поставки №45-А",
    "contract_type": "purchase",
    "counterparty": 'ООО "ТехноПром"',
    "content": "1. СТОРОНЫ\nПоставщик и Покупатель...",
    "amount": 150000000,
    "currency": "UZS",
}


async def _create(client, headers, **overrides):
    resp = await client.post(
        "/api/contracts/", json={**CONTRACT, **overrides}, headers=headers
    )
    assert resp.status_code == 201, resp.text
    return resp.json()


async def test_create_and_get_contract(client, admin_headers):
    created = await _create(client, admin_headers)
    assert created["status"] == "draft"
    assert created["currency"] == "UZS"

    resp = await client.get(f"/api/contracts/{created['id']}", headers=admin_headers)
    assert resp.status_code == 200
    body = resp.json()
    assert body["title"] == CONTRACT["title"]
    assert body["content"] == CONTRACT["content"]


async def test_invalid_contract_type(client, admin_headers):
    resp = await client.post(
        "/api/contracts/",
        json={**CONTRACT, "contract_type": "weird"},
        headers=admin_headers,
    )
    assert resp.status_code == 400


async def test_list_search_and_status_filter(client, admin_headers):
    await _create(client, admin_headers)
    await _create(
        client,
        admin_headers,
        title="Соглашение о неразглашении",
        contract_type="nda",
        counterparty="ИП Иванов",
    )

    all_items = (await client.get("/api/contracts/", headers=admin_headers)).json()
    assert all_items["total"] == 2

    found = (
        await client.get("/api/contracts/?q=поставки", headers=admin_headers)
    ).json()
    assert found["total"] == 1
    assert found["items"][0]["contract_type"] == "purchase"

    drafts = (
        await client.get("/api/contracts/?status_filter=draft", headers=admin_headers)
    ).json()
    assert drafts["total"] == 2
    signed = (
        await client.get(
            "/api/contracts/?status_filter=signed", headers=admin_headers
        )
    ).json()
    assert signed["total"] == 0


async def test_update_creates_version(client, admin_headers):
    created = await _create(client, admin_headers)

    resp = await client.put(
        f"/api/contracts/{created['id']}",
        json={
            "content": "1. СТОРОНЫ\nОбновлённый текст договора...",
            "changes_description": "Правка раздела 1",
        },
        headers=admin_headers,
    )
    assert resp.status_code == 200

    versions = (
        await client.get(
            f"/api/contracts/{created['id']}/versions", headers=admin_headers
        )
    ).json()
    assert [v["version_number"] for v in versions] == [2, 1]
    assert versions[0]["changes_description"] == "Правка раздела 1"

    # обновление без изменения текста не плодит версии
    resp = await client.put(
        f"/api/contracts/{created['id']}",
        json={"counterparty": "Новый контрагент"},
        headers=admin_headers,
    )
    assert resp.status_code == 200
    versions = (
        await client.get(
            f"/api/contracts/{created['id']}/versions", headers=admin_headers
        )
    ).json()
    assert len(versions) == 2


async def test_lawyer_sees_only_own_contracts(client, admin_headers):
    admin_contract = await _create(client, admin_headers)

    resp = await client.post(
        "/api/users/",
        json={"email": "lw@test.uz", "username": "lawyer-user", "password": "Secret1234"},
        headers=admin_headers,
    )
    assert resp.status_code == 201
    login = await client.post(
        "/api/auth/login", json={"email": "lw@test.uz", "password": "Secret1234"}
    )
    lawyer_headers = {"Authorization": f"Bearer {login.json()['access_token']}"}

    own = await client.post(
        "/api/contracts/",
        json={**CONTRACT, "title": "Контракт юриста"},
        headers=lawyer_headers,
    )
    assert own.status_code == 201

    listing = (await client.get("/api/contracts/", headers=lawyer_headers)).json()
    assert listing["total"] == 1
    assert listing["items"][0]["title"] == "Контракт юриста"

    resp = await client.get(
        f"/api/contracts/{admin_contract['id']}", headers=lawyer_headers
    )
    assert resp.status_code == 403

    # админ видит оба
    listing = (await client.get("/api/contracts/", headers=admin_headers)).json()
    assert listing["total"] == 2


async def test_archive_requires_delete_permission(client, admin_headers):
    created = await _create(client, admin_headers)

    resp = await client.post(
        "/api/users/",
        json={
            "email": "head@test.uz",
            "username": "head",
            "password": "Secret1234",
            "role": "head",
        },
        headers=admin_headers,
    )
    assert resp.status_code == 201
    login = await client.post(
        "/api/auth/login", json={"email": "head@test.uz", "password": "Secret1234"}
    )
    head_headers = {"Authorization": f"Bearer {login.json()['access_token']}"}

    # у head нет права delete
    resp = await client.delete(
        f"/api/contracts/{created['id']}", headers=head_headers
    )
    assert resp.status_code == 403

    resp = await client.delete(
        f"/api/contracts/{created['id']}", headers=admin_headers
    )
    assert resp.status_code == 200
    detail = (
        await client.get(f"/api/contracts/{created['id']}", headers=admin_headers)
    ).json()
    assert detail["status"] == "archived"


async def test_upload_docx_parses_text_and_stores_file(client, admin_headers):
    doc = Document()
    doc.add_paragraph("ДОГОВОР АРЕНДЫ")
    doc.add_paragraph("1. Предмет: аренда офисного помещения 120 кв.м.")
    buf = io.BytesIO()
    doc.save(buf)

    resp = await client.post(
        "/api/contracts/upload",
        data={
            "title": "Договор аренды №12-В",
            "contract_type": "lease",
            "counterparty": 'ЗАО "Недвижимость Плюс"',
        },
        files={
            "file": (
                "lease.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers=admin_headers,
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert "ДОГОВОР АРЕНДЫ" in body["content"]
    assert body["file_path"]

    # исходный файл отдаётся стримом (attachment), а не presigned-ссылкой
    resp = await client.get(
        f"/api/contracts/{body['id']}/download", headers=admin_headers
    )
    assert resp.status_code == 200
    assert "attachment" in resp.headers.get("content-disposition", "")
    assert resp.content == buf.getvalue()


async def test_upload_unsupported_extension(client, admin_headers):
    resp = await client.post(
        "/api/contracts/upload",
        data={"title": "X", "contract_type": "other"},
        files={"file": ("virus.exe", b"MZ...", "application/octet-stream")},
        headers=admin_headers,
    )
    assert resp.status_code == 400


async def test_dashboard_metrics_reflect_contracts(client, admin_headers):
    await _create(client, admin_headers)
    metrics = (
        await client.get("/api/dashboard/metrics", headers=admin_headers)
    ).json()
    # черновики не считаются проверенными
    assert metrics["total_reviewed"] == 0
    assert metrics["signed"] == 0
    assert metrics["high_risk"] == 0


async def test_analytics_endpoint(client, admin_headers):
    from datetime import date

    for i, (title, amount, cp) in enumerate(
        [
            ("Аналитика 1", 100, "ООО Альфа"),
            ("Аналитика 2", 300, "ООО Альфа"),
            ("Аналитика 3", 50, "ООО Бета"),
        ]
    ):
        resp = await client.post(
            "/api/contracts/",
            json={
                "title": title,
                "contract_type": "purchase" if i < 2 else "nda",
                "counterparty": cp,
                "amount": amount,
                "content": "Текст",
            },
            headers=admin_headers,
        )
        assert resp.status_code == 201

    data = (
        await client.get("/api/dashboard/analytics?months=3", headers=admin_headers)
    ).json()
    assert len(data["months"]) == 3
    assert data["months"][-1]["month"] == date.today().strftime("%Y-%m")
    assert data["months"][-1]["created"] == 3
    assert data["totals"]["contracts"] == 3
    assert data["risk"]["unscored"] == 3

    types = {t["type"]: t["count"] for t in data["by_type"]}
    assert types == {"purchase": 2, "nda": 1}

    top = data["top_counterparties"]
    assert top[0]["counterparty"] == "ООО Альфа"
    assert top[0]["total_amount"] == 400.0


async def test_export_csv(client, admin_headers):
    resp = await client.post(
        "/api/contracts/",
        json={
            "title": "Экспортный договор",
            "contract_type": "purchase",
            "counterparty": "ООО Экспорт",
            "amount": 1500000.5,
        },
        headers=admin_headers,
    )
    assert resp.status_code == 201

    resp = await client.get("/api/contracts/export/csv", headers=admin_headers)
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith("text/csv")
    assert "attachment" in resp.headers["content-disposition"]

    text = resp.content.decode("utf-8-sig")  # BOM срезается
    lines = text.strip().splitlines()
    assert lines[0].startswith("ID;Название;Тип")
    assert any("Экспортный договор" in line and "1500000,50" in line for line in lines)

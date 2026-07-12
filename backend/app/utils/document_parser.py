"""Извлечение текста контракта из загружаемых файлов (PDF/DOCX/TXT)."""

import io
import zipfile

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt")
MAX_EXTRACTED_CHARS = 5_000_000
MAX_PDF_PAGES = 500


def parse_file(filename: str, data: bytes) -> str:
    name = filename.lower()
    try:
        if name.endswith(".pdf"):
            text = parse_pdf(data)
        elif name.endswith(".docx"):
            text = parse_docx(data)
        elif name.endswith(".txt"):
            text = data.decode("utf-8", errors="strict").strip()
        else:
            raise ValueError(
                f"Неподдерживаемый формат файла. Допустимы: {', '.join(SUPPORTED_EXTENSIONS)}"
            )
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("Не удалось безопасно прочитать документ") from exc
    if len(text) > MAX_EXTRACTED_CHARS:
        raise ValueError("Извлечённый текст документа слишком большой")
    return text


def parse_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ValueError("PDF содержит больше 500 страниц")
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def parse_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def pdf_needs_ocr(filename: str, text: str) -> bool:
    """Return true when a PDF has no meaningful machine-readable text layer."""
    return filename.lower().endswith(".pdf") and len("".join(text.split())) < 200
